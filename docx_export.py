import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding for a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    """Adds a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    
    # Change heading color to deep blue
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(15, 23, 42)
    return heading

def export_to_docx(ddr_data: dict, inspection_data: dict, thermal_data: dict, output_path: str) -> None:
    """
    Renders the compiled DDR data and inline images to a formatted Word document.
    """
    logger.info(f"Generating DOCX report at: {output_path}")
    doc = Document()
    
    # Define styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Cover / Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("DETAILED DIAGNOSTIC REPORT (DDR)")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title.paragraph_format.space_after = Pt(4)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Automated Diagnostics & Thermal Survey")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(79, 70, 229)
    subtitle.paragraph_format.space_after = Pt(24)
    
    # Property Summary Section
    add_heading_styled(doc, "1. Property Summary", level=1)
    
    doc.add_paragraph(ddr_data.get("property_issue_summary", "Not Available"))
    
    # Area-wise Observations Section
    add_heading_styled(doc, "2. Area-wise Observations", level=1)
    
    for area in ddr_data.get("area_wise_observations", []):
        add_heading_styled(doc, f"Area: {area.get('area_name')}", level=2)
        
        # Description
        p_desc = doc.add_paragraph()
        p_desc.add_run(area.get("description", "Not Available"))
        p_desc.paragraph_format.space_after = Pt(12)
        
        # References
        p_refs = doc.add_paragraph()
        p_refs.add_run(f"Inspection Report references: Page(s) {', '.join(map(str, area.get('inspection_page_refs', [])))}\n")
        p_refs.add_run(f"Thermal Survey references: Page(s) {', '.join(map(str, area.get('thermal_page_refs', [])))}")
        p_refs.runs[0].font.size = Pt(9.5)
        p_refs.runs[0].font.italic = True
        p_refs.runs[0].font.color.rgb = RGBColor(100, 116, 139)
        p_refs.paragraph_format.space_after = Pt(12)
        
        # Referenced images
        image_refs = area.get("image_refs", [])
        if image_refs:
            for ref in image_refs:
                source = ref.get("source")
                page_num = ref.get("page")
                img_idx = ref.get("image_index")
                
                target_pdf = inspection_data if source == "inspection" else thermal_data
                
                # Retrieve the PIL Image
                target_page = next((p for p in target_pdf["pages"] if p["page"] == page_num), None)
                if target_page and img_idx < len(target_page["images"]):
                    pil_img = target_page["images"][img_idx]
                    
                    # Convert PIL Image to bytes
                    img_byte_arr = io.BytesIO()
                    pil_img.save(img_byte_arr, format='PNG')
                    img_byte_arr.seek(0)
                    
                    try:
                        # Insert image
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = p_img.add_run()
                        run_img.add_picture(img_byte_arr, width=Inches(3.5))
                        
                        # Add Caption
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_text = f"Figure: {source.capitalize()} Page {page_num} (Image index {img_idx})"
                        run_cap = p_cap.add_run(cap_text)
                        run_cap.font.size = Pt(9)
                        run_cap.font.italic = True
                        run_cap.font.color.rgb = RGBColor(100, 116, 139)
                        p_cap.paragraph_format.space_after = Pt(18)
                    except Exception as e:
                        logger.warning(f"Failed to embed image {source}_p{page_num}_idx{img_idx}: {e}")
                        
    # Root Cause Section
    add_heading_styled(doc, "3. Probable Root Cause Analysis", level=1)
    
    root_causes = ddr_data.get("probable_root_cause", [])
    if root_causes:
        table_rc = doc.add_table(rows=1, cols=2)
        table_rc.style = 'Light Shading Accent 1'
        hdr_cells = table_rc.rows[0].cells
        hdr_cells[0].text = 'Area Name'
        hdr_cells[1].text = 'Root Cause Finding'
        
        for item in root_causes:
            row_cells = table_rc.add_row().cells
            row_cells[0].text = item.get("area_name", "Not Available")
            row_cells[1].text = item.get("cause", "Not Available")
            set_cell_margins(row_cells[0])
            set_cell_margins(row_cells[1])
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Severity Assessment Section
    add_heading_styled(doc, "4. Severity Assessment", level=1)
    
    severities = ddr_data.get("severity_assessment", [])
    if severities:
        table_sv = doc.add_table(rows=1, cols=3)
        table_sv.style = 'Light Shading Accent 1'
        hdr_cells = table_sv.rows[0].cells
        hdr_cells[0].text = 'Area Name'
        hdr_cells[1].text = 'Severity Level'
        hdr_cells[2].text = 'Engineering Reasoning'
        
        for item in severities:
            row_cells = table_sv.add_row().cells
            row_cells[0].text = item.get("area_name", "Not Available")
            row_cells[1].text = item.get("severity", "Not Available")
            row_cells[2].text = item.get("reasoning", "Not Available")
            set_cell_margins(row_cells[0])
            set_cell_margins(row_cells[1])
            set_cell_margins(row_cells[2])

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Recommended Actions
    add_heading_styled(doc, "5. Recommended Actions", level=1)
    
    actions = ddr_data.get("recommended_actions", [])
    for idx, act in enumerate(actions):
        p_act = doc.add_paragraph(style='List Bullet')
        run_title = p_act.add_run(f"{act.get('area_name')}: ")
        run_title.bold = True
        p_act.add_run(act.get("action", "Not Available"))
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Additional Notes
    add_heading_styled(doc, "6. Additional Notes & Guidelines", level=1)
    doc.add_paragraph(ddr_data.get("additional_notes", "Not Available"))
    
    # Missing Information
    add_heading_styled(doc, "7. Missing or Unclear Information", level=1)
    missing_info = ddr_data.get("missing_or_unclear_information", [])
    if missing_info:
        for item in missing_info:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("No missing or unclear information flagged. Report is complete.")
        
    # Save the document
    doc.save(output_path)
